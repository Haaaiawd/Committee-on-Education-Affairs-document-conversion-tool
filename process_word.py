from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os
import sys
from zipfile import BadZipFile
from docx.oxml.ns import qn
from docx.shared import Inches
import shutil
from docx2python import docx2python

def extract_author_number(filename):
    """从文件名中提取作者数字"""
    match = re.search(r'(\d+)', filename)
    return int(match.group(1)) if match else float('inf')

def extract_author_from_filename(filename):
    """从文件名中提取作者名"""
    try:
        # 匹配文件名中852后面的数字，然后后面的2-4个汉字
        author_match = re.search(r'852\d+[^一-龥]*([一-龥]{2,4})', filename)
        if author_match:
            return author_match.group(1).strip()
    except Exception:
        pass
    return None

def has_images_in_doc(doc):
    """
    检查Word文档中是否包含图片
    """
    try:
        for para in doc.paragraphs:
            for run in para.runs:
                try:
                    # 检查内联图片
                    if run._element.findall('.//wp:inline', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}):
                        return True
                    # 检查锚定图片
                    if run._element.findall('.//wp:anchor', namespaces={'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'}):
                        return True
                    # 检查图形对象
                    if run._element.findall('.//w:drawing', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        return True
                except Exception:
                    # 忽略单个图片的检查错误
                    continue
        return False
    except Exception:
        # 如果整个检查过程出错，假设文档包含图片
        return True

def extract_images_from_doc(input_path, temp_dir):
    """
    直接从Word文档关系中提取图片
    """
    try:
        doc = Document(input_path)
        temp_image_files = []
        
        # 获取文档中的所有关系
        rels = doc.part.rels
        
        for rel in rels.values():
            # 检查是否是图片
            if "image" in rel.target_ref:
                try:
                    # 获取图片数据
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # 从目标引用中获取图片扩展名
                    image_ext = os.path.splitext(rel.target_ref)[1]
                    if not image_ext:
                        image_ext = '.png'  # 默认使用png
                    
                    # 保存图片
                    temp_image_path = os.path.join(temp_dir, f"image_{len(temp_image_files)}{image_ext}")
                    with open(temp_image_path, 'wb') as f:
                        f.write(image_data)
                    temp_image_files.append(temp_image_path)
                except Exception as e:
                    print(f"保存图片时出错：{str(e)}")
                    continue
        
        return temp_image_files
    except Exception as e:
        print(f"提取图片时出错：{str(e)}")
        return []

def process_word_file(input_file, output_dir):
    """处理单个Word文件"""
    print(f"DEBUG: 开始处理文件 {input_file}")
    try:
        # 检查文件是否存在
        if not os.path.exists(input_file):
            print(f"× 错误：输入文件 '{input_file}' 不存在")
            return False

        # 创建成功文件文件夹
        success_dir = os.path.join(output_dir, "成功文件")
        no_image_dir = os.path.join(output_dir, "无图片成功文件")
        os.makedirs(success_dir, exist_ok=True)
        os.makedirs(no_image_dir, exist_ok=True)

        # 创建临时文件夹存储图片
        temp_dir = os.path.join(output_dir, "temp_images")
        os.makedirs(temp_dir, exist_ok=True)

        # 打开文档
        try:
            doc = Document(input_file)
        except BadZipFile:
            print(f"× 错误：文件 '{input_file}' 可能已损坏或不是有效的Word文档")
            return False

        # 创建新文档
        new_doc = Document()
        temp_image_files = []
        author_name = ""
        original_title = ""
        title_found = False
        author_added = False  # 添加标志，防止重复添加作者信息

        # 从文件名中提取作者名
        filename = os.path.basename(input_file)
        author_name = extract_author_from_filename(filename)

        # 提取图片
        temp_image_files = extract_images_from_doc(input_file, temp_dir)
        has_images = len(temp_image_files) > 0

        # 处理文档内容
        for para in doc.paragraphs:
            try:
                text = para.text.strip()
                if not text:
                    continue

                # 提取标题（第一个非空段落）
                if not title_found:
                    original_title = text
                    title_para = new_doc.add_paragraph()
                    title_run = title_para.add_run(original_title)
                    title_run.bold = True
                    title_run.font.size = Pt(16)
                    title_run.font.name = '黑体'
                    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    subtitle = '——福州大学先进制造学院与海洋学院关工委2023年"中华魂"（毛泽东伟大精神品格）主题教育征文'
                    subtitle_run = title_para.add_run('\n' + subtitle)
                    subtitle_run.bold = True
                    subtitle_run.font.size = Pt(16)
                    subtitle_run.font.name = '黑体'
                    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                    title_found = True
                    continue

                # 如果从文件名中没有提取到作者名，则尝试从文档内容中提取
                if not author_name and text.startswith('852'):
                    author_match = re.search(r'852\d*[^-]*-([^-\d\W]+)', text)
                    if not author_match:
                        author_match = re.search(r'852\d*[\s-]*([^\d\W]+)', text)
                    if author_match:
                        author_name = author_match.group(1).strip()

                # 添加作者信息（只添加一次）
                if author_name and not author_added and not text.startswith('852'):
                    author_para = new_doc.add_paragraph()
                    author_run = author_para.add_run(f"（先进制造学院与海洋学院关工委通讯员{author_name}）")
                    author_run.bold = True
                    author_run.font.size = Pt(14)
                    author_run.font.name = '宋体'
                    author_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    author_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    author_added = True
                    continue

                # 处理正文（跳过学号行）
                if title_found and not text.startswith('852'):
                    new_para = new_doc.add_paragraph()
                    text_run = new_para.add_run('  ' + text)
                    text_run.font.size = Pt(12)
                    text_run.font.name = '宋体'
                    text_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            except Exception as e:
                print(f"× 处理段落时出错：{str(e)}")
                continue

        # 在文档末尾添加图片
        if temp_image_files:
            new_doc.add_paragraph()  # 添加空行
            for image_path in temp_image_files:
                try:
                    img_para = new_doc.add_paragraph()
                    img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run = img_para.add_run()
                    run.add_picture(image_path, width=Inches(6))
                except Exception as e:
                    print(f"× 添加图片时出错：{str(e)}")
                    continue

        # 保存新文档
        if author_name and original_title:
            new_filename = f"({author_name}){original_title}——福州大学先进制造学院与海洋学院关工委2023年'中华魂'（毛泽东伟大精神品格）主题教育征文.docx"
            # 根据是否有图片选择保存目录
            output_dir_final = success_dir if has_images else no_image_dir
            output_file = os.path.join(output_dir_final, new_filename)
            
            try:
                new_doc.save(output_file)
                if has_images:
                    print(f"✓ 文件处理完成：{new_filename}")
                else:
                    print(f"✓ 文件处理完成（无图片）：{new_filename}")
            except Exception as e:
                print(f"× 保存文件时出错：{str(e)}")
                return False
        else:
            if not author_name:
                print("× 未能提取作者名")
            if not original_title:
                print("× 未能提取标题")
            return False

        # 清理临时文件
        for temp_file in temp_image_files:
            try:
                os.remove(temp_file)
            except:
                pass
        try:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
        except:
            pass

        return True

    except Exception as e:
        print(f"× 处理文件时出现错误：{str(e)}")
        return False

def open_word_doc(input_path):
    """
    安全地打开Word文档
    返回：(doc对象, 错误信息) 元组
    """
    try:
        # 首先尝试直接打开
        doc = Document(input_path)
        return doc, None
    except Exception as e:
        try:
            # 如果直接打开失败，尝试以二进制模式打开
            with open(input_path, 'rb') as f:
                doc = Document(f)
                return doc, None
        except Exception as e:
            error_msg = str(e)
            if "Bad CRC-32" in error_msg:
                return None, "文件中的图片可能已损坏"
            elif "Package not found" in error_msg:
                return None, "文件格式错误或已损坏"
            else:
                return None, f"无法打开文件：{error_msg}"

def process_folder(input_folder, output_folder):
    """处理文件夹中的所有Word文档"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    temp_dir = os.path.join(output_folder, "temp_images")
    
    try:
        for filename in os.listdir(input_folder):
            if filename.startswith('~$') or not filename.endswith('.docx'):
                continue
                
            if hasattr(sys.stdout, 'set_current_file'):
                sys.stdout.set_current_file(filename)
            if hasattr(sys.stderr, 'set_current_file'):
                sys.stderr.set_current_file(filename)
                
            input_path = os.path.join(input_folder, filename)
            print(f"\n处理文件：{filename}")
            
            # 处理文档
            process_word_file(input_path, output_folder)
                    
    except Exception as e:
        print(f"× 处理文件夹时发生错误：{str(e)}")
    finally:
        if os.path.exists(temp_dir):
            try:
                shutil.rmtree(temp_dir)
            except:
                pass
        if hasattr(sys.stdout, 'set_current_file'):
            sys.stdout.set_current_file(None)
        if hasattr(sys.stderr, 'set_current_file'):
            sys.stderr.set_current_file(None)

# 使用示例
if __name__ == "__main__":
    input_folder = r"C:\Users\86159\Desktop\新建文件夹 (2)"
    output_folder = r"C:\Users\86159\Desktop\新建文件夹 (2)"
    
    process_folder(input_folder, output_folder)